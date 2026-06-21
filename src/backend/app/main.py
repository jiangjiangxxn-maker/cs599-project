"""
Career AI Platform - FastAPI Application Entry Point
Multi-agent system for fresh graduate job seekers.
With conversation history persistence and session management.
"""
from __future__ import annotations

import os
from contextlib import asynccontextmanager

from fastapi import FastAPI, HTTPException, UploadFile, File
from fastapi.middleware.cors import CORSMiddleware
from fastapi.responses import StreamingResponse
from pydantic import BaseModel

from app.core.config import config
from app.core.state import OrchestratorState
from app.core.memory import memory_manager
from app.core.tracing import tracing_manager
from app.orchestrator.graph import orchestrator


# ============================================================================
# Request/Response Models
# ============================================================================

class ChatRequest(BaseModel):
    """User chat request"""
    message: str
    session_id: str = ""
    user_id: str = ""


class ChatResponse(BaseModel):
    """Agent response"""
    session_id: str
    response: dict
    result_type: str = "unknown"


# ============================================================================
# Application Lifespan
# ============================================================================

@asynccontextmanager
async def lifespan(app: FastAPI):
    """Initialize services on startup, clean up on shutdown."""
    print(f"[Career AI] Starting {config.app_name} v{config.version}...")
    print(f"[Career AI] Debug mode: {config.debug}")
    print(f"[Career AI] Default LLM: {config.default_agent_llm}")

    await memory_manager.initialize()
    print(f"[Career AI] Memory systems initialized")

    await tracing_manager.initialize()
    print(f"[Career AI] Tracing initialized")

    print(f"[Career AI] LLM Configuration (MoE Routing):")
    print(f"  DeepSeek (reasoning): {'✅ Configured' if config.llm.deepseek_api_key else '❌ Not set'}")
    print(f"    Agent 1(JD分析), Agent 3(岗位匹配), Agent 4(面试模拟)")
    print(f"  Qwen (text polish):   {'✅ Configured' if config.llm.qwen_api_key else '❌ Not set'}")
    print(f"    Agent 0(职业探索), Agent 2(简历优化), Agent 5(学习规划)")

    yield
    print("[Career AI] Shutting down...")


# ============================================================================
# FastAPI Application
# ============================================================================

app = FastAPI(
    title=config.app_name,
    version=config.version,
    description="AI-powered career assistant for fresh graduates",
    lifespan=lifespan,
)

# Frontend directory for serving static files
import os as _os
FRONTEND_DIR = _os.path.join(_os.path.dirname(_os.path.dirname(_os.path.dirname(__file__))), "frontend")
if _os.path.isdir(FRONTEND_DIR):
    print(f"[Career AI] Frontend directory: {FRONTEND_DIR}")

app.add_middleware(
    CORSMiddleware,
    allow_origins=["*"],
    allow_credentials=True,
    allow_methods=["*"],
    allow_headers=["*"],
)


# ============================================================================
# Helper: extract a summary (first user message) from session
# ============================================================================

def _get_session_summary(state: OrchestratorState) -> str:
    """Extract a short summary from the session's first user message."""
    for msg in state.agent_context.conversation_history:
        if msg.get("role") == "user":
            text = msg.get("content", "")
            if len(text) > 30:
                return text[:30] + "..."
            return text
    return "空对话"


# ============================================================================
# Health & Info
# ============================================================================

@app.get("/health")
async def health_check():
    return {
        "status": "ok",
        "app": config.app_name,
        "version": config.version,
        "active_sessions": len(memory_manager.session_memory.list_sessions()),
    }


# ============================================================================
# POST /chat — Main chat endpoint
# ============================================================================

@app.post("/chat", response_model=ChatResponse)
async def chat(request: ChatRequest):
    """
    Main chat endpoint.
    Routes user input to the appropriate agent.
    Saves conversation history to in-memory session.
    """
    if not request.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    # Create or get session
    session_id = request.session_id
    if not session_id or not memory_manager.get_session(session_id):
        session_id = memory_manager.create_session()
        state = memory_manager.get_session(session_id)
    else:
        state = memory_manager.get_session(session_id)

    if not state:
        state = OrchestratorState(session_id=session_id)

    # Save user message
    state.agent_context.conversation_history.append({
        "role": "user",
        "content": request.message,
        "agent_id": -1,
        "agent": "User",
    })

    # Update state with user input
    state.user_input = request.message
    state.pipeline_complete = False

    # Process through orchestrator
    try:
        state = await orchestrator.process(state)
    except Exception as e:
        state.agent_context.errors.append(f"Orchestrator error: {str(e)}")

    # Format response
    response_data = orchestrator.format_response(state)

    # Save agent response
    agent_name = "Unknown"
    if state.agent_context.current_agent_id in orchestrator.agents:
        agent_name = orchestrator.agents[state.agent_context.current_agent_id].name

    # Format the response text for conversation history
    response_text_chunks = _format_result_chunks(response_data)
    response_text = "".join(response_text_chunks) if response_text_chunks else response_data.get("result", {}).get("type", "processed")

    state.agent_context.conversation_history.append({
        "role": "assistant",
        "content": response_text,
        "agent_id": state.agent_context.current_agent_id,
        "agent": agent_name,
        "result": response_data.get("result"),
    })

    # Persist
    memory_manager.update_session(session_id, state)

    result_type = "unknown"
    if "result" in response_data:
        result_type = response_data["result"].get("type", "unknown")

    return ChatResponse(
        session_id=session_id,
        response=response_data,
        result_type=result_type,
    )


# ============================================================================
# SSE Event Helpers
# ============================================================================

def _sse_status(node: str, message: str) -> str:
    """Construct a status SSE event."""
    import json
    return f"data: {json.dumps({'type': 'status', 'node': node, 'message': message}, ensure_ascii=False)}\n\n"


def _sse_message(chunk: str) -> str:
    """Construct a message chunk SSE event."""
    import json
    return f"data: {json.dumps({'type': 'message', 'chunk': chunk}, ensure_ascii=False)}\n\n"


def _sse_done(session_id: str, next_agent: str = "", result_type: str = "", response: dict = None) -> str:
    """Construct a done SSE event."""
    import json
    payload = {
        'type': 'done',
        'session_id': session_id,
        'next_suggested_agent': next_agent,
        'result_type': result_type,
    }
    if response:
        payload['response'] = response
    return f"data: {json.dumps(payload, ensure_ascii=False)}\n\n"


def _format_result_chunks(response_data: dict) -> list[str]:
    """
    Split structured result into text chunks for typewriter effect.
    Returns a list of string chunks.
    """
    chunks = []
    result = response_data.get("result", {})
    result_type = result.get("type", "")

    if result_type == "career_feasibility_report":
        chunks.append("根据你的背景，以下是岗位可行性分析：\n\n")
        for p in result.get("positions", []):
            chunks.append(f"**{p.get('name', '')}** (匹配度 {p.get('match_score', 0)*100:.0f}%)\n")
            if p.get("daily_tasks"):
                chunks.append(f"  日常工作: {', '.join(p['daily_tasks'])}\n")
            if p.get("skill_barriers"):
                chunks.append(f"  技能门槛: {', '.join(p['skill_barriers'])}\n")
            chunks.append("\n")
        if result.get("skill_gaps"):
            chunks.append("**技能差距:**\n")
            for g in result["skill_gaps"]:
                chunks.append(f"  - {g.get('skill', '')} 需要从 {g.get('from', '入门')} 提升到 {g.get('to', '熟练')}\n")
        if result.get("confidence_score"):
            chunks.append(f"\n置信度: {result['confidence_score']*100:.0f}%\n")

    elif result_type == "jd_analysis_report":
        chunks.append(f"**岗位分析** (难度: {result.get('overall_difficulty', 'medium')})\n\n")
        for r in result.get("requirements", []):
            badge = "[硬性要求]" if r.get("category") == "hard_filter" else "[加分项]" if r.get("category") == "nice_to_have" else "[工作职责]"
            chunks.append(f"{badge} {r.get('original', '')}\n")
            chunks.append(f"  -> {r.get('translated', '')}\n\n")
        if result.get("gaps"):
            chunks.append("**技能差距:**\n")
            for g in result["gaps"]:
                chunks.append(f"  - {g.get('skill', '')}: {g.get('action', '')}\n")
        if result.get("suggestion"):
            chunks.append(f"\n{result['suggestion']}\n")

    elif result_type == "resume_optimization":
        chunks.append("**简历优化建议:**\n\n")
        for c in result.get("comparisons", []):
            chunks.append(f"~~{c.get('original', '')}~~\n")
            chunks.append(f"  -> {c.get('optimized', '')}\n\n")
        if result.get("ats_score"):
            chunks.append(f"ATS 评分: {result['ats_score']}/100\n")
        if result.get("suggestions"):
            chunks.append("\n**建议:**\n")
            for s in result["suggestions"]:
                chunks.append(f"  - {s}\n")

    elif result_type == "interview_simulation":
        mode_label = {"technical": "技术面", "hr": "HR面", "project": "项目深挖"}.get(result.get("mode", ""), "面试")
        is_interviewing = result.get("is_interviewing", False)
        if is_interviewing:
            # Active interview: show latest question/feedback
            qa_list = result.get("qa", [])
            if qa_list:
                last_qa = qa_list[-1]
                has_answer = bool(last_qa.get("user_answer") or last_qa.get("feedback"))
                if has_answer:
                    # User just answered → show feedback + next question or follow-up
                    chunks.append(f"**反馈：** {last_qa.get('feedback', '')}\n")
                    if last_qa.get("score"):
                        chunks.append(f"评分：{last_qa['score']}/10\n")
                    if last_qa.get("question"):
                        chunks.append(f"\n**{last_qa['question']}**\n")
                else:
                    # Phase A: first question, no answer yet
                    chunks.append(f"**{last_qa.get('question', '面试模拟进行中...')}**\n")
            else:
                chunks.append("**面试模拟进行中...**\n")
        else:
            # Interview completed: show full evaluation report
            chunks.append(f"**面试模拟 ({mode_label})**\n\n")
            for q in result.get("qa", []):
                chunks.append(f"**Q:** {q.get('question', '')}\n")
                if q.get("score"):
                    chunks.append(f"  评分: {q['score']}/10\n")
                if q.get("feedback"):
                    chunks.append(f"  {q['feedback']}\n")
                chunks.append("\n")
            if result.get("overall_score"):
                chunks.append(f"**总分: {result['overall_score']}/10**\n")
            if result.get("strengths"):
                chunks.append(f"\n优势: {', '.join(result['strengths'])}\n")
            if result.get("weaknesses"):
                chunks.append(f"薄弱点: {', '.join(result['weaknesses'])}\n")

    elif result_type == "learning_plan":
        chunks.append(f"**学习规划** ({result.get('duration_weeks', 8)}周)\n\n")
        for w in result.get("weekly_plan", []):
            chunks.append(f"第{w.get('week', '')}周: {w.get('topic', '')} ({w.get('hours', 0)}小时)\n")
            if w.get("objectives"):
                chunks.append(f"  目标: {', '.join(w['objectives'])}\n")
            chunks.append("\n")
        if result.get("projects"):
            chunks.append("**推荐项目:**\n")
            for p in result["projects"]:
                chunks.append(f"  - {p.get('name', '')} ({p.get('difficulty', '')})\n")
                if p.get("tech_stack"):
                    chunks.append(f"    技术栈: {', '.join(p['tech_stack'])}\n")

    elif result_type == "job_matching":
        chunks.append("**岗位匹配结果:**\n\n")
        for p in result.get("matched_positions", []):
            chunks.append(f"**{p.get('company', '')}** - {p.get('position', '')} (匹配度 {p.get('match_score', 0)*100:.0f}%)\n")
            if p.get("days_left"):
                chunks.append(f"  剩余 {p['days_left']} 天\n")
            chunks.append("\n")
        if result.get("strategy"):
            chunks.append(f"**策略:** {result['strategy']}\n")

    else:
        chunks.append("分析完成。\n")

    return chunks


@app.post("/chat/stream")
async def chat_stream(request: ChatRequest):
    """
    SSE streaming endpoint with structured event protocol.
    Events: status, message, done.
    """
    import json
    import asyncio

    if not request.message.strip():
        raise HTTPException(status_code=400, detail="Message cannot be empty")

    # Create or get session
    session_id = request.session_id
    if not session_id or not memory_manager.get_session(session_id):
        session_id = memory_manager.create_session()
        state = memory_manager.get_session(session_id)
    else:
        state = memory_manager.get_session(session_id)

    if not state:
        state = OrchestratorState(session_id=session_id)
    else:
        # Use centralized reset to prevent state leakage between turns
        state.reset_for_new_request()

    # Save user message
    state.agent_context.conversation_history.append({
        "role": "user",
        "content": request.message,
        "agent_id": -1,
        "agent": "User",
    })

    state.user_input = request.message

    async def event_generator():
        nonlocal state
        try:
            # Step 1: Semantic cache check
            yield _sse_status("Orchestrator", "正在检查语义缓存...")
            await asyncio.sleep(0.05)

            # Step 2: Intent routing
            yield _sse_status("Orchestrator", "正在通过语义路由分析您的求职意图...")
            state = await orchestrator.route_request(state)
            agent_id = state.agent_context.current_agent_id
            agent_obj = orchestrator.agents.get(agent_id)
            agent_display = f"Agent{agent_id}_{agent_obj.name}" if agent_obj else f"Agent{agent_id}"

            yield _sse_status("Orchestrator", f"意图识别完毕，唤醒【{agent_obj.name if agent_obj else agent_display}】")
            await asyncio.sleep(0.05)

            # Step 3: Run agent
            yield _sse_status(agent_display, "正在分析中...")

            try:
                state = await orchestrator.run_agent(agent_id, state)
            except Exception as e:
                state.agent_context.errors.append(f"Orchestrator error: {str(e)}")

            # Step 4: Self-reflection check
            yield _sse_status("SelfReflection", "正在检查是否需要跨 Agent 唤醒...")
            state = await orchestrator._post_process(state)
            await asyncio.sleep(0.05)

            # Step 5: Format response and extract response text for streaming
            response_data = orchestrator.format_response(state)

            result = response_data.get("result", {})
            result_type = result.get("type", "unknown")
            ctx = state.agent_context

            # Extract response text to stream: conversation_history last assistant msg
            stream_text = ""
            for msg in reversed(ctx.conversation_history):
                if msg.get("role") == "assistant" and msg.get("agent_id", -1) >= 0:
                    stream_text = msg.get("content", "")
                    break
            if not stream_text:
                stream_text = "\n".join(_format_result_chunks(response_data))

            # Save agent response — Agent 4 manages its own history, NEVER auto-append
            if ctx.current_agent_id != 4:
                agent_name_str = "Unknown"
                if ctx.current_agent_id in orchestrator.agents:
                    agent_name_str = orchestrator.agents[ctx.current_agent_id].name
                # Use stream_text which contains the formatted response
                saved_content = stream_text if stream_text else result_type
                ctx.conversation_history.append({
                    "role": "assistant",
                    "content": saved_content,
                    "agent_id": ctx.current_agent_id,
                    "agent": agent_name_str,
                    "result": result,
                })

            memory_manager.update_session(session_id, state)

            # Step 6: Stream response text token by token (true typewriter)
            if stream_text:
                yield _sse_status(agent_display, "正在生成回复...")
                await asyncio.sleep(0.05)

                import re as _sse_re
                tokens = _sse_re.findall(r'[\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]|[^\u4e00-\u9fff\u3000-\u303f\uff00-\uffef]+', stream_text)
                for token in tokens:
                    yield _sse_message(token)
                    await asyncio.sleep(0.02)

            # Step 7: Done event
            next_agent = ""
            if ctx.current_agent_id == 4 and ctx.learning_plan:
                next_agent = "Agent5_LearningPlanner"

            yield _sse_done(session_id, next_agent=next_agent, result_type=result_type, response=response_data)

        except Exception as _sse_err:
            import traceback as _sse_tb
            _sse_tb.print_exc()
            yield _sse_message(f"\n⚠️ 系统异常：{str(_sse_err)}")
        finally:
            yield "data: [DONE]\n\n"

    return StreamingResponse(
        event_generator(),
        media_type="text/event-stream",
        headers={
            "Cache-Control": "no-cache",
            "Connection": "keep-alive",
            "X-Accel-Buffering": "no",
        },
    )


# ============================================================================
# GET /sessions — List all sessions with summaries
# ============================================================================

@app.get("/sessions")
async def list_sessions():
    """List all sessions with summary and message count. Filters out empty chats."""
    sessions = []
    for sid in memory_manager.session_memory.list_sessions():
        state = memory_manager.get_session(sid)
        if not state:
            continue
        msg_count = len(state.agent_context.conversation_history)
        # Skip empty sessions (no user messages)
        if msg_count < 2:
            continue
        sessions.append({
            "session_id": sid,
            "summary": _get_session_summary(state),
            "message_count": msg_count,
            "error_count": len(state.agent_context.errors),
            "pipeline_complete": state.pipeline_complete,
        })
    # Sort newest first (most recent message first)
    sessions.sort(key=lambda s: s["message_count"], reverse=True)
    return {"sessions": sessions, "total": len(sessions)}


# ============================================================================
# GET /session/{session_id} — Session details
# ============================================================================

@app.get("/session/{session_id}")
async def get_session(session_id: str):
    """Get session details."""
    state = memory_manager.get_session(session_id)
    if not state:
        raise HTTPException(status_code=404, detail="Session not found")

    return {
        "session_id": session_id,
        "current_agent": state.agent_context.current_agent_id,
        "pipeline_complete": state.pipeline_complete,
        "error_count": len(state.agent_context.errors),
        "conversation_length": len(state.agent_context.conversation_history),
        "summary": _get_session_summary(state),
    }


# ============================================================================
# GET /session/{session_id}/messages — Full conversation
# ============================================================================

@app.get("/session/{session_id}/messages")
async def get_session_messages(session_id: str):
    """Get full conversation history for a session."""
    state = memory_manager.get_session(session_id)
    if not state:
        raise HTTPException(status_code=404, detail="Session not found")

    messages = []
    for msg in state.agent_context.conversation_history:
        messages.append({
            "role": msg.get("role", "unknown"),
            "content": msg.get("content", ""),
            "agent": msg.get("agent", ""),
            "result": msg.get("result"),
        })

    return {
        "session_id": session_id,
        "messages": messages,
        "total": len(messages),
    }


# ============================================================================
# DELETE /session/{session_id}
# ============================================================================

@app.delete("/session/{session_id}")
async def delete_session(session_id: str):
    """Delete a session."""
    state = memory_manager.get_session(session_id)
    if not state:
        raise HTTPException(status_code=404, detail="Session not found")

    memory_manager.session_memory.delete_session(session_id)
    return {"status": "deleted", "session_id": session_id}


# ============================================================================
# GET /agents
# ============================================================================

@app.get("/agents")
async def list_agents():
    """List all available agents."""
    agents_info = []
    for agent_id, agent in orchestrator.agents.items():
        agents_info.append({
            "id": agent_id,
            "name": agent.name,
            "llm_provider": agent.llm_provider,
        })
    return {"agents": agents_info, "total": len(agents_info)}


# ============================================================================
# POST /upload-resume — Upload resume file and extract text
# ============================================================================

@app.post("/upload-resume")
async def upload_resume(file: UploadFile = File(...)):
    """
    Upload a resume file (PDF/DOCX/TXT) and extract text content.
    Returns the extracted text for further processing.
    """
    # Validate file type
    allowed_types = {
        "application/pdf": "pdf",
        "application/vnd.openxmlformats-officedocument.wordprocessingml.document": "docx",
        "text/plain": "txt",
    }
    file_ext = file.filename.rsplit(".", 1)[-1].lower() if "." in file.filename else ""
    if file_ext not in ("pdf", "docx", "txt"):
        raise HTTPException(status_code=400, detail="Only PDF, DOCX, TXT files are supported")

    # Read file content
    content = await file.read()
    if len(content) > 5 * 1024 * 1024:  # 5MB limit
        raise HTTPException(status_code=400, detail="File size exceeds 5MB limit")

    text = ""
    try:
        if file_ext == "txt":
            text = content.decode("utf-8", errors="ignore")
        elif file_ext == "pdf":
            import io
            from pdfminer.high_level import extract_text_to_fp
            from pdfminer.layout import LAParams
            output = io.StringIO()
            laparams = LAParams(
                all_texts=True,
                detect_vertical=True,
                word_margin=0.1,
                char_margin=2.0,
                line_margin=0.5,
                boxes_flow=0.5,
            )
            extract_text_to_fp(io.BytesIO(content), output, laparams=laparams, output_type="text")
            text = output.getvalue()
        elif file_ext == "docx":
            import io
            from docx import Document
            doc = Document(io.BytesIO(content))
            # Extract all paragraphs
            for para in doc.paragraphs:
                if para.text.strip():
                    text += para.text + "\n"
            # Extract all tables (important for structured resumes)
            for table in doc.tables:
                for row in table.rows:
                    row_text = " | ".join(cell.text.strip() for cell in row.cells)
                    if row_text.strip():
                        text += row_text + "\n"
                text += "\n"
    except Exception as e:
        raise HTTPException(status_code=500, detail=f"Failed to parse file: {str(e)}")

    if not text.strip():
        raise HTTPException(status_code=400, detail="Could not extract text from file")

    return {
        "filename": file.filename,
        "file_type": file_ext,
        "text": text.strip(),
        "char_count": len(text.strip()),
    }


# ============================================================================
# Catch-all: serve frontend static files for non-API routes
# This must be the LAST route to avoid hijacking API paths
# ============================================================================

from fastapi.responses import FileResponse as _FileResponse


@app.get("/{full_path:path}")
async def serve_frontend(full_path: str):
    """Serve frontend static files or index.html as SPA fallback."""
    file_path = os.path.join(FRONTEND_DIR, full_path)
    if os.path.isfile(file_path):
        return _FileResponse(file_path)
    index_path = os.path.join(FRONTEND_DIR, "index.html")
    if os.path.isfile(index_path):
        return _FileResponse(index_path)
    raise HTTPException(status_code=404, detail="Not found")


# ============================================================================
# Entry point
# ============================================================================

if __name__ == "__main__":
    import uvicorn
    uvicorn.run(
        "app.main:app",
        host=config.host,
        port=config.port,
        reload=config.debug,
    )
